import io
from fastapi import UploadFile, HTTPException
from backend.models.schemas import ParsedDocument


async def parse_document(file: UploadFile) -> ParsedDocument:
    """
    Extract plain text from an uploaded resume file.
    Supports PDF, DOCX, and plain TXT formats.
    """
    content = await file.read()
    filename = file.filename or ""

    if filename.endswith(".pdf"):
        return _parse_pdf(content)
    elif filename.endswith(".docx"):
        return _parse_docx(content)
    elif filename.endswith(".txt") or file.content_type == "text/plain":
        return ParsedDocument(
            text=content.decode("utf-8", errors="ignore"),
            source_type="text",
        )
    else:
        raise HTTPException(
            status_code=415,
            detail="Unsupported file type. Please upload a PDF, DOCX, or TXT file.",
        )


def _parse_pdf(content: bytes) -> ParsedDocument:
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text_parts.append(extracted)

        return ParsedDocument(
            text="\n".join(text_parts),
            page_count=page_count,
            source_type="pdf",
        )
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not parse PDF: {str(e)}")


def _parse_docx(content: bytes) -> ParsedDocument:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        return ParsedDocument(
            text="\n".join(paragraphs),
            source_type="docx",
        )
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not parse DOCX: {str(e)}")
